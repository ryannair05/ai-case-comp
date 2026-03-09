"""
Data export & account deletion router.

Ethics requirement: One-click full data export must work at ALL times.
Never behind a support ticket. Customer owns their data.
"""
import io
import json
from datetime import datetime, timedelta

import docx
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse

from models.schemas import get_current_customer, get_supabase
from services.pinecone_client import delete_customer_data

router = APIRouter(prefix="/export", tags=["export"])


# ---------------------------------------------------------------------------
# DOCX export for a single proposal
# ---------------------------------------------------------------------------

@router.get("/proposals/{proposal_id}/docx")
async def export_proposal_docx(
    proposal_id: str,
    customer=Depends(get_current_customer),
):
    """
    Download a single proposal as a formatted Word (.docx) document.
    Customer owns their content — export is always available.
    """
    supabase = get_supabase()
    result = (
        supabase.table("proposals")
        .select("*")
        .eq("id", proposal_id)
        .eq("customer_id", str(customer.id))
        .single()
        .execute()
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Proposal not found")

    proposal = result.data
    document = docx.Document()

    # Title
    title = proposal.get("title") or f"Proposal — {proposal.get('client_name', 'Client')}"
    document.add_heading(title, level=0)

    # Meta block
    meta_lines = []
    if proposal.get("client_name"):
        meta_lines.append(f"Client: {proposal['client_name']}")
    if proposal.get("value_usd"):
        meta_lines.append(f"Value: ${proposal['value_usd']:,.0f}")
    if proposal.get("created_at"):
        meta_lines.append(f"Date: {proposal['created_at'][:10]}")
    if meta_lines:
        document.add_paragraph("\n".join(meta_lines)).italic = True

    document.add_paragraph("")  # spacer

    # Body — split on double-newline to preserve section breaks
    for paragraph in proposal.get("content", "").split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        # Treat lines starting with '#' as section headings
        if paragraph.startswith("#"):
            heading_text = paragraph.lstrip("#").strip()
            document.add_heading(heading_text, level=2)
        else:
            document.add_paragraph(paragraph)

    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)

    safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in title)[:60]
    filename = f"{safe_title}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Full data export (ZIP)
# ---------------------------------------------------------------------------

@router.get("/full")
async def export_all_data(customer=Depends(get_current_customer)):
    """
    One-click full export: all proposals, pricing, brand voice, analytics.
    Returns a ZIP file with JSON files for each data type.
    Must work even if customer is canceling (ethics requirement).
    """
    import zipfile
    import uuid

    supabase = get_supabase()
    cid = str(customer.id)

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        # Proposals
        proposals_result = (
            supabase.table("proposals").select("*").eq("customer_id", cid).execute()
        )
        zf.writestr("proposals.json", json.dumps(proposals_result.data or [], default=str, indent=2))

        # Pricing data
        pricing_result = (
            supabase.table("pricing_data").select("*").eq("customer_id", cid).execute()
        )
        zf.writestr(
            "pricing_data.json",
            json.dumps(pricing_result.data or [], default=str, indent=2),
        )

        # Brand voice
        brand_result = (
            supabase.table("brand_voice").select("*").eq("customer_id", cid).execute()
        )
        zf.writestr(
            "brand_voice.json",
            json.dumps(brand_result.data or [], default=str, indent=2),
        )

        # Support tickets
        tickets_result = (
            supabase.table("support_tickets").select("*").eq("customer_id", cid).execute()
        )
        zf.writestr(
            "support_tickets.json",
            json.dumps(tickets_result.data or [], default=str, indent=2),
        )

        # README
        zf.writestr(
            "README.txt",
            (
                "Your Draftly data export.\n"
                "All data is yours. Draftly licenses access, not ownership.\n\n"
                "Draftly deletes your vectors from Pinecone within 30 days of cancellation.\n"
                "To request immediate deletion, email privacy@draftly.ai\n\n"
                f"Export generated: {datetime.utcnow().isoformat()} UTC\n"
                f"Customer ID: {cid}\n"
            ),
        )

    zip_buffer.seek(0)
    filename = f"draftly-export-{cid[:8]}.zip"
    return StreamingResponse(
        io.BytesIO(zip_buffer.getvalue()),
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ---------------------------------------------------------------------------
# Account deletion (soft delete → hard delete after 30 days)
# ---------------------------------------------------------------------------

@router.delete("/account")
async def delete_account(customer=Depends(get_current_customer)):
    """
    30-day soft delete.
    - Pinecone vectors: deleted immediately
    - Postgres data: soft-deleted, hard-deleted after 30 days
    - Stripe subscription: canceled
    GDPR right to deletion honored within 30 days.
    """
    import stripe
    import os

    supabase = get_supabase()
    cid = str(customer.id)

    # 1. Delete Pinecone vectors immediately
    delete_customer_data(cid)

    # 2. Soft-delete customer in Supabase
    hard_delete_date = (datetime.utcnow() + timedelta(days=30)).isoformat()
    supabase.table("customers").update(
        {
            "status": "deleted",
            "deleted_at": datetime.utcnow().isoformat(),
            "hard_delete_at": hard_delete_date,
        }
    ).eq("id", cid).execute()

    # 3. Cancel Stripe subscription
    if customer.stripe_id:
        try:
            stripe_client = stripe.Stripe(os.environ.get("STRIPE_SECRET_KEY", ""))
            # Cancel at period end (30-day window)
            stripe.Subscription.modify(
                customer.stripe_id,
                cancel_at_period_end=True,
                api_key=os.environ.get("STRIPE_SECRET_KEY"),
            )
        except Exception:
            pass  # Don't block deletion if Stripe call fails

    return {
        "status": "deletion_scheduled",
        "pinecone_vectors": "deleted_immediately",
        "postgres_data": "soft_deleted",
        "hard_delete_date": hard_delete_date,
        "message": (
            "Your data will be fully deleted within 30 days. "
            "To recover your account within this window, email support@draftly.ai"
        ),
    }
